from docx import Document
from docx.shared import Inches

document=Document()
#profile picture 
document.add_picture(
    'Denis M.JPG',width=Inches(2.0))
#name phone number and email details 

name=input('What is your name? ')
phone_number=input('What is your phone_number? ')
email=input('What is your email? ')

document.add_paragraph(
    name +'| '+ phone_number + '| '+ email)
#about me

document.add_heading('About me')

document.add_paragraph(
    input('Tell about yourself?'))
#work experience
document.add_heading ('Work Experience')
p=document.add_paragraph()
company=input('Enter Company ')
from_date=input('From Date ')
to_date=input('To Date ')

p.add_run(company+ ' ').bold=True
p.add_run(from_date +' _' +to_date +'\n').italic=True

experience_details=input(
    'Describe your experience at' + company)
p.add_run(experience_details)

 
p=document.add_paragraph(skill)
p.style='List Bullet'

#more experiences 


# footer
section=document .section[0]
footer=section.footer
p=footer.paragraph[0]
p.text='CV generated using denismucyocode and Intuit    QuickBook'


document.save('cv.docx')